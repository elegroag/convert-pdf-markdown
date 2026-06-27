"""Helper to build DOCX files for tests."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches


def build_simple_docx(path: Path | None = None) -> Path | BytesIO:
    """Build a DOCX with heading, paragraph, list, and table."""
    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("Hello **world** paragraph.")
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    if path is not None:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_docx_with_image(path: Path) -> Path:
    """Build a DOCX with a minimal embedded PNG image."""
    doc = Document()
    doc.add_heading("With Image", level=1)
    doc.add_paragraph("Text before image.")

    # Minimal 1x1 PNG
    png_bytes = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx"
        b"\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    doc.add_picture(BytesIO(png_bytes), width=Inches(1.0))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
